import streamlit as st
import pandas as pd
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from webdriver_manager.core.os_manager import ChromeType
from bs4 import BeautifulSoup
from collections import Counter
import time
from io import BytesIO
from docx import Document

# --- CẤU HÌNH TRANG ---
st.set_page_config(page_title="SEO Content Researcher", layout="wide")

# CSS tùy chỉnh cho đẹp
st.markdown("""
<style>
    .main {background-color: #f4f6f9;}
    h1 {color: #2c3e50; font-family: 'Helvetica', sans-serif;}
    .stButton>button {width: 100%; border-radius: 5px; font-weight: bold;}
    .stTextArea textarea {font-family: monospace;}
</style>
""", unsafe_allow_html=True)

# --- QUẢN LÝ TRẠNG THÁI (SESSION STATE) ---
if 'results' not in st.session_state:
    st.session_state['results'] = None
if 'is_analyzed' not in st.session_state:
    st.session_state['is_analyzed'] = False

# --- HÀM CÀI ĐẶT DRIVER ---
@st.cache_resource
def get_driver():
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
    
    service = Service(ChromeDriverManager(chrome_type=ChromeType.CHROMIUM).install())
    driver = webdriver.Chrome(service=service, options=chrome_options)
    return driver

# --- HÀM LÀM SẠCH HTML (CLEANER) ---
def clean_html(soup):
    # 1. Xóa các thẻ kỹ thuật & thẻ điều hướng chắc chắn không chứa content
    for tag in soup(['script', 'style', 'header', 'footer', 'nav', 'aside', 'noscript', 'form', 'iframe']):
        tag.decompose()
    
    # 2. Xóa theo Class/ID rác (Sidebar, Menu, Related, Comments)
    # Các từ khóa thường gặp trong class của web
    garbage_keywords = [
        'sidebar', 'widget', 'menu', 'nav', 'comment', 'share', 'social', 
        'popup', 'modal', 'cookie', 'related', 'author-box', 'breadcrumb', 'footer'
    ]
    
    for tag in soup.find_all(True):
        # Kiểm tra class và id của thẻ
        check_list = (tag.get('class') or []) + ([tag.get('id')] if tag.get('id') else [])
        check_str = " ".join(check_list).lower()
        
        if any(kw in check_str for kw in garbage_keywords):
            tag.decompose()
            
    return soup

# --- GIAO DIỆN CHÍNH ---
st.title("🔎 SEO Content Researcher & Outline Generator")

if not st.session_state['is_analyzed']:
    # Màn hình nhập liệu
    with st.container():
        st.info("💡 Tool hỗ trợ quét nội dung JS, chặn bot (An Cường, v.v). Tối đa 5 URL.")
        urls_input = st.text_area("👉 Dán danh sách URL (Mỗi dòng 1 link):", height=200)
        
        if st.button("🚀 BẮT ĐẦU PHÂN TÍCH", type="primary"):
            if not urls_input.strip():
                st.warning("Vui lòng nhập ít nhất 1 URL!")
            else:
                url_list = [x.strip() for x in urls_input.split('\n') if x.strip()][:5] # Giới hạn 5 URL
                
                # --- BẮT ĐẦU QUÉT ---
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                all_data = [] # Dữ liệu thô để xuất Excel
                outline_corpus = [] # Chứa H2 để tìm điểm chung
                display_text_full = "" # Biến chứa text hiển thị trên web
                
                try:
                    driver = get_driver()
                    
                    for i, url in enumerate(url_list):
                        status_text.text(f"⏳ Đang xử lý ({i+1}/{len(url_list)}): {url}")
                        
                        page_data = {'URL': url, 'Title': '', 'Meta Desc': '', 'Headings': []}
                        
                        try:
                            driver.get(url)
                            time.sleep(3) # Chờ JS load
                            soup = BeautifulSoup(driver.page_source, 'html.parser')
                            
                            # 1. Lấy SEO Title & Meta (Trước khi clean)
                            if soup.title:
                                page_data['Title'] = soup.title.get_text(strip=True)
                            
                            meta = soup.find('meta', attrs={'name': 'description'}) or soup.find('meta', attrs={'property': 'og:description'})
                            if meta:
                                page_data['Meta Desc'] = meta.get('content', '').strip()

                            # 2. Clean HTML
                            soup = clean_html(soup)
                            
                            # 3. Lấy Heading
                            headings = soup.find_all(['h1', 'h2', 'h3'])
                            heading_list_text = []
                            
                            # Format text hiển thị
                            display_text_full += f"URL: {url}\n"
                            display_text_full += f"TITLE: {page_data['Title']}\n"
                            display_text_full += f"META: {page_data['Meta Desc']}\n"
                            display_text_full += "STRUCTURE:\n"

                            for tag in headings:
                                txt = tag.get_text(strip=True)
                                if txt:
                                    tag_name = tag.name.upper()
                                    heading_list_text.append(f"[{tag_name}] {txt}")
                                    display_text_full += f"- [{tag_name}] {txt}\n"
                                    
                                    if tag.name == 'h2':
                                        outline_corpus.append(txt)

                            display_text_full += "\n" + "="*50 + "\n\n"
                            page_data['Headings'] = "\n".join(heading_list_text) # Lưu dạng chuỗi để cho vào Excel
                            
                            all_data.append(page_data)

                        except Exception as e:
                            st.error(f"Lỗi khi đọc {url}: {e}")
                        
                        progress_bar.progress((i + 1) / len(url_list))
                    
                    # Xử lý Logic Outline Recommend
                    recommend_outline = []
                    if outline_corpus:
                        # Làm sạch text để so sánh
                        normalized_h2 = [h.lower().replace('là gì','').replace('như thế nào','').strip() for h in outline_corpus]
                        # Đếm tần suất
                        counter = Counter(normalized_h2)
                        
                        # Logic: Lấy các ý xuất hiện > 1 lần, nếu ít quá thì lấy top 10
                        most_common = counter.most_common(15)
                        recommend_text = "GỢI Ý OUTLINE (Dựa trên tần suất xuất hiện):\n"
                        for topic, count in most_common:
                            original_text = next((h for h in outline_corpus if h.lower().replace('là gì','').replace('như thế nào','').strip() == topic), topic.title())
                            note = f"(x{count} web nhắc đến)" if count > 1 else ""
                            recommend_text += f"- {original_text} {note}\n"
                            recommend_outline.append(original_text)
                    else:
                        recommend_text = "Không đủ dữ liệu H2 để đề xuất outline."

                    # LƯU VÀO SESSION STATE
                    st.session_state['results'] = {
                        'all_data': all_data,
                        'display_text': display_text_full,
                        'recommend_text': recommend_text,
                        'recommend_list': recommend_outline
                    }
                    st.session_state['is_analyzed'] = True
                    st.rerun() # Load lại trang để chuyển sang màn hình kết quả

                except Exception as e:
                    st.error(f"Lỗi khởi động Driver: {e}")

else:
    # --- MÀN HÌNH KẾT QUẢ ---
    res = st.session_state['results']
    
    # Nút Trở lại (Reset)
    if st.button("⬅️ TRỞ LẠI (F5 để research mới)"):
        st.session_state['results'] = None
        st.session_state['is_analyzed'] = False
        st.rerun()

    col1, col2 = st.columns(2)
    
    # 1. Hiển thị Outline Recommend
    with col1:
        st.subheader("💡 Outline Đề Xuất")
        st.text_area("Copy Outline:", value=res['recommend_text'], height=400)
    
    # 2. Hiển thị Chi tiết Research
    with col2:
        st.subheader("📝 Dữ liệu chi tiết (Raw)")
        st.text_area("Toàn bộ Title, Meta, Heading:", value=res['display_text'], height=400)

    st.divider()
    st.subheader("📂 Xuất Dữ Liệu")
    
    c1, c2 = st.columns(2)
    
    # --- NÚT XUẤT WORD (DOCX) ---
    doc = Document()
    doc.add_heading('BÁO CÁO NGHIÊN CỨU SEO', 0)
    
    doc.add_heading('PHẦN 1: OUTLINE ĐỀ XUẤT', level=1)
    for line in res['recommend_list']:
        doc.add_paragraph(f"- {line}", style='List Bullet')
        
    doc.add_heading('PHẦN 2: CHI TIẾT ĐỐI THỦ', level=1)
    for item in res['all_data']:
        doc.add_heading(item['URL'], level=2)
        doc.add_paragraph(f"SEO Title: {item['Title']}")
        doc.add_paragraph(f"Meta Desc: {item['Meta Desc']}")
        doc.add_paragraph("Headings:")
        doc.add_paragraph(item['Headings'])
        doc.add_paragraph("-" * 20)

    buffer_doc = BytesIO()
    doc.save(buffer_doc)
    buffer_doc.seek(0)
    
    with c1:
        st.download_button(
            label="📄 Tải file Word (.docx)",
            data=buffer_doc,
            file_name="SEO_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )

    # --- NÚT XUẤT EXCEL (2 SHEETS) ---
    buffer_xls = BytesIO()
    with pd.ExcelWriter(buffer_xls, engine='xlsxwriter') as writer:
        # Sheet 1: Outline Research
        # Chuẩn bị data cho Sheet 1
        sheet1_data = []
        for idx, item in enumerate(res['all_data']):
            sheet1_data.append({
                'No.': idx + 1,
                'URL': item['URL'],
                'Title SEO': item['Title'],
                'Headings (H2-H3)': item['Headings'], # Đã format xuống dòng trong text
                'Meta description': item['Meta Desc']
            })
        df1 = pd.DataFrame(sheet1_data)
        df1.to_excel(writer, sheet_name='Outline Research', index=False)
        
        # Format cột Headings cho dễ đọc (Wrap text)
        workbook = writer.book
        worksheet1 = writer.sheets['Outline Research']
        format_wrap = workbook.add_format({'text_wrap': True, 'valign': 'top'})
        worksheet1.set_column('D:D', 50, format_wrap) # Cột Headings rộng ra
        worksheet1.set_column('B:B', 30) # URL
        worksheet1.set_column('C:C', 30) # Title
        worksheet1.set_column('E:E', 40, format_wrap) # Meta

        # Sheet 2: Outline Recommend
        df2 = pd.DataFrame(res['recommend_list'], columns=['Recommended H2'])
        df2.to_excel(writer, sheet_name='Outline Recommend', index=False)
        
        # Format Sheet 2
        worksheet2 = writer.sheets['Outline Recommend']
        worksheet2.set_column('A:A', 60)

    buffer_xls.seek(0)
    
    with c2:
        st.download_button(
            label="📊 Tải file Excel (2 Sheets)",
            data=buffer_xls,
            file_name="SEO_Research_Data.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            type="primary"
        )